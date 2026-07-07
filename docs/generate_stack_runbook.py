#!/usr/bin/env python3
"""
CloudLens Stack Deployment Runbook generator (AWS).

Builds CloudLens_Stack_Deployment_Runbook.docx using the same visual language as
generate_runbook.py:
  - Keysight gold logo, AWS squid-ink (#232F3E) section headers
  - AWS-navy header rows on tables (white text)
  - Courier New code blocks
  - Calibri body, professional sans-serif
  - Cover page + static (LibreOffice-safe) TOC + page numbers

Audience: customer DevOps, CTO/procurement, training. Covers the three
deployment paths (Bash one-liner / Terraform stack module / AWS Console Launch
Stack) plus prerequisites, verification, troubleshooting and cleanup.

Run:
    python3 docs/generate_stack_runbook.py
Output:
    docs/CloudLens_Stack_Deployment_Runbook.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# -----------------------------------------------------------------------------
# Brand palette (matches generate_runbook.py exactly)
# -----------------------------------------------------------------------------
AWS_ORANGE = RGBColor(0xFF, 0x99, 0x00)
AWS_NAVY = RGBColor(0x23, 0x2F, 0x3E)
AWS_BLUE = RGBColor(0x14, 0x6E, 0xB4)
KEYSIGHT_GOLD = RGBColor(0xD4, 0xAF, 0x37)
KEYSIGHT_NAVY = RGBColor(0x1B, 0x2A, 0x4A)
TEXT_DARK = RGBColor(0x1B, 0x2A, 0x4A)
TEXT_MUTED = RGBColor(0x55, 0x65, 0x75)
CALLOUT_BG = "FFF3E0"
CODE_BG = "F2F2F2"
TABLE_HEADER_BG = "232F3E"
TABLE_ALT_ROW_BG = "F5F7FA"
ACCENT_HEX = "FF9900"

# -----------------------------------------------------------------------------
# Paths
# -----------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent
OUTPUT_DOCX = SCRIPT_DIR / "CloudLens_Stack_Deployment_Runbook.docx"


# -----------------------------------------------------------------------------
# XML helpers
# -----------------------------------------------------------------------------
def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = "BFBFBF", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_MUTED


# Static, pre-evaluated TOC: kept in sync with build_* calls in main().
TOC_ENTRIES: list[tuple[int, str]] = [
    (1, "1. Executive Summary"),
    (1, "2. The Three Deployment Paths"),
    (1, "3. Prerequisites"),
    (1, "4. Path A: Bash One-Liner"),
    (2, "4.1  What the script does (phase by phase)"),
    (2, "4.2  Flags and overrides"),
    (1, "5. Path B: Terraform Stack Module"),
    (2, "5.1  Workflow"),
    (2, "5.2  What gets created"),
    (1, "6. Path C: AWS Console (Launch Stack)"),
    (1, "7. Verification Checklist"),
    (1, "8. Troubleshooting Reference"),
    (1, "9. Cleanup / Decommission"),
    (1, "10. Appendix: File Paths and Quick Links"),
]


def _add_toc_entry(doc: Document, level: int, label: str) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    if level >= 2:
        pf.left_indent = Inches(0.35)
    tab_stops = pf.tab_stops
    try:
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    except Exception:
        tab_stops.add_tab_stop(Inches(6.3))
    label_run = p.add_run(label)
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(12) if level == 1 else Pt(10.5)
    label_run.bold = (level == 1)
    label_run.font.color.rgb = AWS_BLUE if level == 1 else TEXT_DARK
    p.add_run("\t")


# -----------------------------------------------------------------------------
# Style configuration
# -----------------------------------------------------------------------------
def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_DARK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = AWS_NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = AWS_BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = KEYSIGHT_NAVY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True


# -----------------------------------------------------------------------------
# Reusable element builders
# -----------------------------------------------------------------------------
def add_paragraph(doc: Document, text: str, *, bold: bool = False,
                  italic: bool = False, color: RGBColor | None = None,
                  size: int = 11, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODE_BG)
    p_pr.append(shd)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = KEYSIGHT_NAVY


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _shade_cell(cell, CALLOUT_BG)
    _set_cell_borders(cell, color=ACCENT_HEX, size="6")
    label_p = cell.paragraphs[0]
    label_p.paragraph_format.space_after = Pt(2)
    run = label_p.add_run(label)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = AWS_NAVY
    body_p = cell.add_paragraph()
    run2 = body_p.add_run(body)
    run2.font.name = "Calibri"
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = TEXT_DARK
    doc.add_paragraph()


def add_checkbox_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.1)
        box = p.add_run("☐  ")
        box.font.name = "Segoe UI Symbol"
        box.font.size = Pt(13)
        box.font.color.rgb = AWS_BLUE
        text = p.add_run(item)
        text.font.name = "Calibri"
        text.font.size = Pt(11)
        text.font.color.rgb = TEXT_DARK


def add_numbered(doc: Document, steps: list[str]) -> None:
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK


def add_styled_table(doc: Document, headers: list[str], rows: list[list[str]],
                     col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _shade_cell(cell, TABLE_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            cell = row_cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                _shade_cell(cell, TABLE_ALT_ROW_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_DARK
    if col_widths:
        for row in table.rows:
            for c_idx, width_in in enumerate(col_widths):
                row.cells[c_idx].width = Inches(width_in)
    doc.add_paragraph()


# -----------------------------------------------------------------------------
# Header / Footer
# -----------------------------------------------------------------------------
def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.different_first_page_header_footer = True

    header = section.header
    h_p = header.paragraphs[0]
    h_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = h_p.add_run("CloudLens Stack Deployment Runbook (AWS)")
    h_run.font.name = "Calibri"
    h_run.font.size = Pt(9)
    h_run.font.color.rgb = AWS_BLUE
    h_run.italic = True

    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = f_p.add_run("Keysight Technologies   |   ")
    left.font.name = "Calibri"
    left.font.size = Pt(9)
    left.font.color.rgb = TEXT_MUTED
    _add_page_number(f_p)


# -----------------------------------------------------------------------------
# Cover page
# -----------------------------------------------------------------------------
def build_cover_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_p.add_run("KEYSIGHT")
    logo_run.font.name = "Calibri"
    logo_run.font.size = Pt(20)
    logo_run.bold = True
    logo_run.font.color.rgb = KEYSIGHT_GOLD

    tagline_p = doc.add_paragraph()
    tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tagline_p.add_run("TECHNOLOGIES")
    tag_run.font.name = "Calibri"
    tag_run.font.size = Pt(10)
    tag_run.font.color.rgb = KEYSIGHT_NAVY
    tag_run.bold = True

    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CloudLens Stack")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(36)
    title_run.bold = True
    title_run.font.color.rgb = AWS_NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Deployment Runbook for AWS")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(26)
    sub_run.font.color.rgb = AWS_ORANGE

    doc.add_paragraph()
    doc.add_paragraph()

    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag_p.add_run(
        "vController, KVO, vPB and sensors. End to end. One command."
    )
    tag_run.font.name = "Calibri"
    tag_run.font.size = Pt(14)
    tag_run.italic = True
    tag_run.font.color.rgb = TEXT_MUTED

    for _ in range(8):
        doc.add_paragraph()

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    for row in meta_table.rows:
        for cell in row.cells:
            cell.width = Inches(2.0)
    labels = [("Version", "v1.0"), ("Date", "June 2026")]
    for r_idx, (label, value) in enumerate(labels):
        l_cell = meta_table.rows[r_idx].cells[0]
        v_cell = meta_table.rows[r_idx].cells[1]
        _shade_cell(l_cell, TABLE_HEADER_BG)
        _shade_cell(v_cell, TABLE_ALT_ROW_BG)
        lp = l_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        lr.font.name = "Calibri"
        lr.font.size = Pt(11)
        vp = v_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run(value)
        vr.bold = True
        vr.font.color.rgb = KEYSIGHT_NAVY
        vr.font.name = "Calibri"
        vr.font.size = Pt(11)

    for _ in range(2):
        doc.add_paragraph()

    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = foot_p.add_run("Keysight Technologies   |   Network Visibility Solutions")
    foot_run.font.name = "Calibri"
    foot_run.font.size = Pt(10)
    foot_run.italic = True
    foot_run.font.color.rgb = TEXT_MUTED

    doc.add_page_break()


# -----------------------------------------------------------------------------
# Sections
# -----------------------------------------------------------------------------
def build_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    for level, label in TOC_ENTRIES:
        _add_toc_entry(doc, level, label)
    doc.add_page_break()


def build_executive_summary(doc: Document) -> None:
    doc.add_heading("1. Executive Summary", level=1)
    add_paragraph(
        doc,
        "This runbook shows how to deploy the full CloudLens stack on AWS in one "
        "shot: vController (formerly CLMS), KVO (Keysight Vision One), the "
        "Virtual Packet Broker (vPB), and the OS sensors that connect tagged EC2 "
        "instances back to vController. Three paths are documented; pick the one "
        "that matches how your team works. All three produce the same AWS "
        "resources.",
    )

    doc.add_heading("Audience", level=2)
    add_bullets(doc, [
        "Customer DevOps and platform engineers who own the AWS account.",
        "Keysight Sales Engineers running customer POCs or demos.",
        "Customer CTO / procurement / training teams who need a printable reference.",
    ])

    doc.add_heading("What gets deployed", level=2)
    add_styled_table(
        doc,
        headers=["Component", "AWS resources", "Required"],
        rows=[
            ["vController (CLMS)",
             "1 EC2 (t3.xlarge), 1 Elastic IP, security group, ENI",
             "Yes"],
            ["KVO (Keysight Vision One)",
             "1 EC2 (c5.2xlarge), 1 Elastic IP, security group, ENI",
             "Optional"],
            ["Virtual Packet Broker (vPB)",
             "1 EC2 (t3.xlarge), 1 Elastic IP, security group, ENI",
             "Optional"],
            ["Collector SVM",
             "Auto-deployed by KVO Zone Tapping when a Cloud Config commits",
             "Auto"],
            ["Shared networking",
             "VPC 10.99.0.0/16 with mgmt, data and tool subnets, IGW, route table",
             "Yes"],
        ],
        col_widths=[1.9, 3.5, 1.2],
    )

    doc.add_heading("Time and cost", level=2)
    add_styled_table(
        doc,
        headers=["Phase", "Wall-clock time"],
        rows=[
            ["CloudFormation stack CREATE_COMPLETE", "about 5 min"],
            ["vController + KVO init (EULA reachable)", "about 15 min"],
            ["vPB reachable on port 9022", "10 to 15 min"],
            ["Sensors on first 10 instances (optional)", "5 to 8 min"],
            ["Total typical (stack + config + sensors)", "30 to 45 min"],
        ],
        col_widths=[3.6, 2.4],
    )
    add_paragraph(
        doc,
        "Indicative AWS compute cost (us-east-1, June 2026 on-demand list): "
        "vController t3.xlarge and vPB t3.xlarge about USD 120/month each, KVO "
        "c5.2xlarge about USD 250/month, plus small EBS and Elastic IP fees. "
        "Stopping the instances stops compute charges. Marketplace software fees "
        "are billed separately by Keysight.",
        italic=True, size=10, color=TEXT_MUTED,
    )
    doc.add_page_break()


def build_paths_overview(doc: Document) -> None:
    doc.add_heading("2. The Three Deployment Paths", level=1)
    add_paragraph(
        doc,
        "All three paths use the same Marketplace AMIs, deploy the same EC2 "
        "instances, and chain vController first, then KVO and vPB. Pick the one "
        "that matches your operating model. You can switch between them later "
        "without losing state, because each path produces standard AWS resources "
        "in a CloudFormation stack or Terraform state.",
    )
    add_styled_table(
        doc,
        headers=["Path", "Best for", "Tools needed", "Time to first paste"],
        rows=[
            ["A. Bash one-liner",
             "First-time trial, customer demo, CloudShell users",
             "Browser + AWS CLI (CloudShell has both)",
             "10 seconds"],
            ["B. Terraform stack module",
             "IaC pipelines, repeatable customer envs, GitOps",
             "Terraform 1.5+ and AWS CLI logged in",
             "2 minutes"],
            ["C. AWS Console (Launch Stack)",
             "Click-through customers, governance-controlled accounts",
             "Browser only",
             "1 minute"],
        ],
        col_widths=[1.7, 2.5, 1.9, 1.3],
    )
    add_callout(
        doc, "Why three?",
        "Customers do not all work the same way. A platform team wants Terraform. "
        "A network engineer trialling the kit wants a paste. A CTO walking through "
        "with procurement wants a button. Three paths, identical outcome.",
    )
    doc.add_page_break()


def build_prerequisites(doc: Document) -> None:
    doc.add_heading("3. Prerequisites", level=1)
    add_paragraph(
        doc,
        "Confirm each item before kicking off the deployment. The script stops "
        "early with a clear message if any of the first three are missing, so you "
        "do not waste time mid-deploy.",
    )
    add_checkbox_list(doc, [
        "AWS account with rights to create a CloudFormation stack, VPC, EC2, and "
        "IAM resources (or AdministratorAccess for the SE-led path).",
        "AWS CLI v2 installed and logged in (SSO or access keys). In CloudShell "
        "this is preinstalled and pre-authenticated.",
        "Subscribed to the three Keysight Marketplace AMIs (Vision One, CloudLens "
        "Manager, Virtual Packet Broker). This is a one-time interactive step per "
        "account and cannot be automated.",
        "An EC2 key pair created in the target region (us-east-1 by default).",
        "vCPU quota headroom: c5.2xlarge (KVO, 8 vCPU) and t3.xlarge (vController "
        "and vPB, 4 vCPU each) in your region.",
        "Open egress on the workload subnets to reach vController on TCP/443 "
        "(only needed once sensors are deployed).",
    ])
    add_callout(
        doc, "Region note",
        "Default region is us-east-1 and the AMIs are region-locked. To deploy "
        "elsewhere, look up the equivalent Marketplace AMI IDs first: aws ec2 "
        "describe-images --owners aws-marketplace --filters "
        "Name=name,Values=*CloudLens*Manager* --region <region>. See "
        "docs/OPERATIONS.md section 9.",
    )
    doc.add_page_break()


def build_path_a(doc: Document) -> None:
    doc.add_heading("4. Path A: Bash One-Liner", level=1)
    add_paragraph(
        doc,
        "The fastest way to a working stack. One paste in AWS CloudShell or any "
        "local terminal with the AWS CLI; the script drives CloudFormation and "
        "then the sensor playbook.",
    )
    add_code_block(
        doc,
        "curl -sSL "
        "https://raw.githubusercontent.com/Keysight-Tech/"
        "cloudlens-ansible-aws/main/deploy/deploy-stack.sh | bash",
    )
    add_paragraph(doc, "Prefer to inspect before running? Download then execute:")
    add_code_block(
        doc,
        "curl -sSL -o deploy-stack.sh \\\n"
        "  https://raw.githubusercontent.com/Keysight-Tech/"
        "cloudlens-ansible-aws/main/deploy/deploy-stack.sh\n"
        "less deploy-stack.sh\n"
        "bash deploy-stack.sh --key-pair my-keypair",
    )

    doc.add_heading("4.1  What the script does (phase by phase)", level=2)
    add_paragraph(
        doc,
        "Each phase prints a banner with its name, so if anything fails you know "
        "exactly where you are.",
    )
    add_numbered(doc, [
        "Banner and environment detection: prints the version header and detects "
        "CloudShell vs local terminal so prompts adapt accordingly.",
        "Pre-flight checks: confirms the AWS CLI is installed and you are logged "
        "in, and warns if the target region is missing quota for the c5 and t3 "
        "families.",
        "Customer input: prompts for stack name (default cloudlens-stack), region "
        "(default us-east-1), EC2 key pair, and admin ingress CIDR. All defaults "
        "can be overridden by flags (see 4.2).",
        "Marketplace check: reminds you to subscribe to the three Keysight AMIs "
        "if the account is not already subscribed. This one step is interactive "
        "on the Marketplace pages.",
        "Stack deploy: submits deploy/cloudformation/stack.yaml. Waits for "
        "CREATE_COMPLETE and reads the Outputs (EIPs and private IPs).",
        "Wait for init: polls the vController and KVO HTTPS endpoints until the "
        "UI is reachable (typically 15 minutes). Prints a one-line progress bar.",
        "Manual project key step: the script pauses and tells you to open the "
        "vController UI, accept the EULA, create a project, and paste the project "
        "key back into the terminal. This is the one step that cannot be "
        "automated end to end today.",
        "Sensor chain (optional): hands off to quickstart.sh, which deploys "
        "sensors to every EC2 instance tagged cloudlens=yes. Prints a final "
        "summary and writes cloudlens-deploy-summary.txt with all IPs and creds.",
    ])
    add_callout(
        doc, "Output you can keep",
        "Every run writes 'cloudlens-deploy-stack.log' (raw stdout/stderr) and "
        "'cloudlens-deploy-summary.txt' (stack name, instance IDs, EIPs, default "
        "creds). Save both to your customer ticket or runbook archive.",
    )

    doc.add_heading("4.2  Flags and overrides", level=2)
    add_styled_table(
        doc,
        headers=["Flag", "Effect"],
        rows=[
            ["--dry-run",
             "Walk through every prompt and print the aws commands without "
             "touching AWS. Use this to preview the run."],
            ["--stack-name NAME",
             "Override the default stack name (cloudlens-stack)."],
            ["--region REGION",
             "Override the default region (us-east-1). Pass a region where you "
             "have subscribed to the Marketplace AMIs."],
            ["--key-pair NAME",
             "EC2 key pair for OS-level SSH access. Required."],
            ["--admin-cidr CIDR",
             "Source CIDR allowed to reach SSH, vPB SSH (9022), HTTPS, and VXLAN. "
             "Narrow this from the 0.0.0.0/0 default."],
            ["--no-kvo / --no-vpb",
             "Skip KVO or vPB. vController-only deploys complete fastest."],
            ["--no-sensors",
             "Stop after the stack phase. Run quickstart.sh later yourself."],
            ["-h | --help",
             "Print the help banner with all flags and a phase summary."],
        ],
        col_widths=[2.0, 4.6],
    )
    doc.add_page_break()


def build_path_b(doc: Document) -> None:
    doc.add_heading("5. Path B: Terraform Stack Module", level=1)
    add_paragraph(
        doc,
        "The stack module wraps the clms, kvo and vpb child modules behind one "
        "tfvars file. One 'terraform apply' provisions the VPC, subnets, "
        "security groups, ENIs, Elastic IPs and all three EC2 instances. The "
        "sensors are still deployed separately via Ansible because they touch "
        "customer instances, not AWS resources.",
    )

    doc.add_heading("5.1  Workflow", level=2)
    add_numbered(doc, [
        "Clone the repository and change into the stack module directory.",
        "Copy terraform.tfvars.example to terraform.tfvars and fill in region, "
        "key_pair_name and admin_cidr. Toggle deploy_kvo / deploy_vpb if you "
        "only want vController.",
        "Run terraform init to download the aws provider.",
        "Run terraform plan to review the resources Terraform will create. "
        "Confirm against your account's policies.",
        "Run terraform apply. About 5 to 10 minutes later, the stack is up. The "
        "outputs print the vController, KVO and vPB URLs and the default creds.",
        "Open the vController UI, accept the EULA, create your project, copy the "
        "project key, and run quickstart.sh to deploy sensors.",
    ])
    add_code_block(
        doc,
        "git clone "
        "https://github.com/Keysight-Tech/cloudlens-ansible-aws.git\n"
        "cd cloudlens-ansible-aws/deploy/terraform/stack\n"
        "cp terraform.tfvars.example terraform.tfvars\n"
        "$EDITOR terraform.tfvars   # set region, key_pair_name, admin_cidr\n"
        "terraform init\n"
        "terraform plan\n"
        "terraform apply",
    )

    doc.add_heading("5.2  What gets created", level=2)
    add_styled_table(
        doc,
        headers=["Resource", "Count", "Notes"],
        rows=[
            ["VPC", "1", "10.99.0.0/16 (override via tfvars)"],
            ["Subnets", "3", "mgmt 10.99.1.0/24, data 10.99.11.0/24, tool 10.99.12.0/24"],
            ["Internet gateway + route table", "1", "Public egress for mgmt subnet"],
            ["Elastic IPs", "1 to 3", "One per deployed component"],
            ["Security groups", "1 to 3", "Scoped to the admin CIDR per role"],
            ["EC2 instances", "1 to 3", "vController t3.xlarge, KVO c5.2xlarge, vPB t3.xlarge"],
        ],
        col_widths=[2.2, 0.9, 3.9],
    )
    add_callout(
        doc, "Existing VPC?",
        "The child modules (deploy/terraform/clms, kvo, vpb) can each target an "
        "existing VPC and subnet if you pass vpc_id and subnet_id instead of "
        "letting the stack create the shared VPC. Use this when the customer "
        "mandates bring-your-own-network.",
    )
    doc.add_page_break()


def build_path_c(doc: Document) -> None:
    doc.add_heading("6. Path C: AWS Console (Launch Stack)", level=1)
    add_paragraph(
        doc,
        "If you prefer to deploy from the Console, the same CloudFormation "
        "template is exposed through the Launch Stack button on the README and "
        "the landing page at keysight-tech.github.io/cloudlens-ansible-aws. The "
        "button opens the CloudFormation quick-create screen with the template "
        "URL pre-loaded from this repository.",
    )
    add_numbered(doc, [
        "Click Launch Stack. Sign in to the AWS Console if prompted.",
        "The quick-create form opens with the stack template already loaded from "
        "the repo's raw GitHub URL.",
        "Fill the parameters: EC2 key pair, admin ingress CIDR, and whether to "
        "deploy KVO and vPB (both default to yes). Instance types are fixed and "
        "should be left at their defaults.",
        "Tick the IAM capability acknowledgement and click Create Stack.",
        "Wait about 5 minutes for CREATE_COMPLETE, then read the Outputs tab for "
        "the vController, KVO and vPB URLs and IPs.",
        "Open the vController UI, accept the EULA, create a project, copy the key, "
        "and run quickstart.sh from CloudShell to deploy sensors.",
    ])
    add_callout(
        doc, "Marketplace subscription first",
        "The stack launches Marketplace AMIs. If the account has not subscribed "
        "to Keysight Vision One, CloudLens Manager, and CloudLens Virtual Packet "
        "Broker, the instances fail to launch. Subscribe once per account on the "
        "Marketplace listing pages before clicking Launch Stack.",
    )
    add_callout(
        doc, "No screenshots, by design",
        "The AWS Console shifts every few months. Rather than ship stale "
        "screenshots, this guide describes the click path in words. The buttons "
        "and field names match the Console at June 2026.",
    )
    doc.add_page_break()


def build_verification(doc: Document) -> None:
    doc.add_heading("7. Verification Checklist", level=1)
    add_paragraph(
        doc,
        "Walk this list after every deploy. If any item fails, jump to section 8 "
        "for the matching troubleshooting entry.",
    )
    add_checkbox_list(doc, [
        "CloudFormation stack shows CREATE_COMPLETE and the Outputs tab lists the "
        "vController, KVO and vPB EIPs.",
        "vController UI reachable on https://<vcontroller-eip>/. The Keysight "
        "CloudLens login screen loads, not a connection error.",
        "Logged in to vController with admin / Cl0udLens@dm!n and completed the "
        "forced password change.",
        "Project created in vController under Settings > Projects, and the API "
        "key copied into a safe place.",
        "vPB reachable on SSH port 9022 after 10 to 15 minutes: ssh -p 9022 "
        "admin@<vpb-eip> (password ixia, then change it).",
        "KVO > Inventory > Devices shows vController and vPB as CONNECTED after "
        "adoption.",
        "Sensors visible in vController > Sensors within 8 minutes of running "
        "quickstart.sh, each showing Connected.",
    ])
    add_callout(
        doc, "Printable",
        "This page is intentionally a single checklist so it can be printed and "
        "ticked off during a customer handover or installation walkthrough.",
    )
    doc.add_page_break()


def build_troubleshooting(doc: Document) -> None:
    doc.add_heading("8. Troubleshooting Reference", level=1)
    add_paragraph(
        doc,
        "Organised by failure mode. The full reference lives at "
        "docs/OPERATIONS.md and docs/TROUBLESHOOTING.md in the repository.",
    )

    doc.add_heading("Marketplace and instance types", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely cause", "Fix"],
        rows=[
            ["Stack creates but instances fail to launch",
             "Not subscribed to the Marketplace AMIs",
             "Subscribe once per account to Vision One, CloudLens Manager, and vPB"],
            ["UnsupportedOperation on create",
             "Wrong instance type for an AMI",
             "KVO must be c5.2xlarge; vController and vPB must be t3.xlarge"],
        ],
        col_widths=[2.0, 2.0, 3.0],
    )

    doc.add_heading("IAM and quota", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely cause", "Fix"],
        rows=[
            ["AccessDenied on stack create",
             "Principal cannot create IAM resources",
             "Deploy with CAPABILITY_NAMED_IAM (CFN) or AdministratorAccess"],
            ["InsufficientInstanceCapacity / quota error",
             "Region quota too low for c5 or t3",
             "Request a quota increase, or deploy in a region with headroom"],
        ],
        col_widths=[2.0, 2.0, 3.0],
    )

    doc.add_heading("Network", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely cause", "Fix"],
        rows=[
            ["vController UI not reachable after 20 minutes",
             "Security group blocks 443 from your client IP",
             "Narrow but include your IP in the admin ingress CIDR"],
            ["vPB SSH times out on 9022",
             "Security group missing TCP/9022, or KCOS still booting",
             "Add the SG rule; wait 10 to 15 min after the instance is running"],
            ["Sensor cannot reach vController",
             "Egress blocked from instance subnet to vController on 443",
             "Open egress on the workload subnet or peer the VPCs"],
        ],
        col_widths=[2.0, 2.0, 3.0],
    )

    doc.add_heading("Timing", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely cause", "Fix"],
        rows=[
            ["vController UI returns 502 / 503 just after deploy",
             "Services still initialising in sequence",
             "Wait the full 15 minutes; the script's poll loop handles this"],
            ["KVO EULA blocks the API",
             "EULA not accepted in a browser",
             "Open https://<kvo-eip>/ and click Agree; the CLI cannot do it"],
        ],
        col_widths=[2.0, 2.0, 3.0],
    )
    doc.add_page_break()


def build_cleanup(doc: Document) -> None:
    doc.add_heading("9. Cleanup / Decommission", level=1)
    add_paragraph(
        doc,
        "Two ways to undo a stack deploy. Pick the one that matches how you "
        "deployed it.",
    )

    doc.add_heading("Terraform deploy", level=2)
    add_paragraph(
        doc,
        "If you used Path B (Terraform), use terraform destroy. This removes "
        "exactly what Terraform created and leaves anything else in the account "
        "untouched.",
    )
    add_code_block(doc, "cd deploy/terraform/stack\nterraform destroy")

    doc.add_heading("Bash or Console deploy", level=2)
    add_paragraph(
        doc,
        "If you used Path A or Path C, delete the CloudFormation stack. This "
        "removes the VPC, instances, Elastic IPs and security groups the stack "
        "created. Double-check the stack name before pressing enter.",
    )
    add_code_block(
        doc,
        "aws cloudformation delete-stack --stack-name cloudlens-stack "
        "--region us-east-1\n"
        "aws cloudformation wait stack-delete-complete "
        "--stack-name cloudlens-stack --region us-east-1",
    )

    doc.add_heading("Sensors on workload instances", level=2)
    add_paragraph(
        doc,
        "Sensors are not removed by terraform destroy or delete-stack, because "
        "they live on the customer's workload instances (usually a different "
        "account or VPC). To remove them:",
    )
    add_bullets(doc, [
        "Linux: docker stop cloudlens-agent && docker rm cloudlens-agent",
        "Windows: Stop-Service CloudLensAgent; sc.exe delete CloudLensAgent",
        "Or: untag the instances (remove cloudlens=yes) and rerun cleanup.yaml",
    ])
    add_callout(
        doc, "Marketplace billing",
        "Deleting the stack stops the Keysight Marketplace subscription billing "
        "for those instances at the next hourly tick. Compute billing stops "
        "immediately. Allow up to 24 hours for the Marketplace line to disappear "
        "from your AWS bill.",
    )
    doc.add_page_break()


def build_appendix(doc: Document) -> None:
    doc.add_heading("10. Appendix: File Paths and Quick Links", level=1)

    doc.add_heading("Repository file paths", level=2)
    add_styled_table(
        doc,
        headers=["File", "Purpose"],
        rows=[
            ["deploy/deploy-stack.sh", "Path A entry point (Bash one-liner)"],
            ["deploy/cloudformation/stack.yaml", "Stack template (Paths A and C)"],
            ["deploy/terraform/stack/", "Path B entry point (Terraform module)"],
            ["deploy/terraform/{clms,kvo,vpb}/", "Per-component Terraform child modules"],
            ["docs/CloudLens_Stack_Deployment_Runbook.pdf", "This document (PDF)"],
            ["docs/CloudLens_Stack_Deployment_Runbook.docx", "This document (Word)"],
            ["quickstart.sh", "Sensor deployment chain (called by deploy-stack.sh)"],
            ["demo/setup-aws-visibility-demo.sh", "Full demo orchestrator (SE lab)"],
        ],
        col_widths=[3.2, 3.8],
    )

    doc.add_heading("Command reference", level=2)
    add_styled_table(
        doc,
        headers=["Action", "Command"],
        rows=[
            ["Run full stack (recommended)",
             "curl -sSL https://raw.githubusercontent.com/Keysight-Tech/"
             "cloudlens-ansible-aws/main/deploy/deploy-stack.sh | bash"],
            ["Run full stack (dry-run)",
             "bash deploy/deploy-stack.sh --dry-run"],
            ["Run vController only",
             "bash deploy/deploy-stack.sh --no-kvo --no-vpb"],
            ["Terraform apply",
             "cd deploy/terraform/stack && terraform apply"],
            ["Terraform destroy",
             "cd deploy/terraform/stack && terraform destroy"],
            ["Delete the stack",
             "aws cloudformation delete-stack --stack-name cloudlens-stack --region us-east-1"],
        ],
        col_widths=[2.2, 4.8],
    )

    doc.add_heading("Useful links", level=2)
    add_styled_table(
        doc,
        headers=["Resource", "Location"],
        rows=[
            ["GitHub repository",
             "https://github.com/Keysight-Tech/cloudlens-ansible-aws"],
            ["Landing page",
             "https://keysight-tech.github.io/cloudlens-ansible-aws/"],
            ["Sensor runbook (companion document)",
             "docs/CloudLens_Ansible_AWS_Customer_Runbook.pdf"],
            ["Operations guide",
             "docs/OPERATIONS.md"],
            ["Troubleshooting reference",
             "docs/TROUBLESHOOTING.md"],
            ["Issues / support",
             "https://github.com/Keysight-Tech/cloudlens-ansible-aws/issues"],
        ],
        col_widths=[2.6, 4.4],
    )
    add_paragraph(
        doc,
        "Support email subject (for the Keysight account team): "
        "'CloudLens Stack AWS: <customer name>, <brief issue>'.",
        italic=True, size=10, color=TEXT_MUTED,
    )


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------
def main() -> int:
    print(f"[info] writing {OUTPUT_DOCX}")

    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)

    build_cover_page(doc)
    build_toc(doc)
    build_executive_summary(doc)
    build_paths_overview(doc)
    build_prerequisites(doc)
    build_path_a(doc)
    build_path_b(doc)
    build_path_c(doc)
    build_verification(doc)
    build_troubleshooting(doc)
    build_cleanup(doc)
    build_appendix(doc)

    doc.save(OUTPUT_DOCX)
    size_kb = OUTPUT_DOCX.stat().st_size / 1024
    print(f"[ok]  wrote {OUTPUT_DOCX.name} ({size_kb:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
