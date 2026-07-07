#!/usr/bin/env python3
"""
CloudLens Ansible for AWS: Customer Runbook generator.

Builds CloudLens_Ansible_AWS_Customer_Runbook.docx, mirroring the polished
README in an executive/printable format. Visual style:
  - Keysight gold logo, AWS squid-ink (#232F3E) section headers
  - AWS-navy header rows on tables (white text)
  - Courier New code blocks
  - Calibri body, professional sans-serif
  - Cover page + static (LibreOffice-safe) TOC + page numbers

Embeds the SVG assets in docs/assets/ by rasterising them to PNG via cairosvg.
If cairosvg is unavailable, falls back to a text caption so the document still
builds end to end.

Run:
    python3 docs/generate_runbook.py
Output:
    docs/CloudLens_Ansible_AWS_Customer_Runbook.docx
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# -----------------------------------------------------------------------------
# Brand palette (AWS + Keysight)
# -----------------------------------------------------------------------------
AWS_ORANGE = RGBColor(0xFF, 0x99, 0x00)
AWS_NAVY = RGBColor(0x23, 0x2F, 0x3E)   # squid ink
AWS_BLUE = RGBColor(0x14, 0x6E, 0xB4)   # link blue
KEYSIGHT_GOLD = RGBColor(0xD4, 0xAF, 0x37)
KEYSIGHT_NAVY = RGBColor(0x1B, 0x2A, 0x4A)
SUCCESS_GREEN = RGBColor(0x22, 0xC5, 0x5E)
TEXT_DARK = RGBColor(0x1B, 0x2A, 0x4A)
TEXT_MUTED = RGBColor(0x55, 0x65, 0x75)
CALLOUT_BG = "FFF3E0"           # light orange
CODE_BG = "F2F2F2"
TABLE_HEADER_BG = "232F3E"      # AWS navy
TABLE_ALT_ROW_BG = "F5F7FA"
ACCENT_HEX = "FF9900"           # AWS orange, callout border

# -----------------------------------------------------------------------------
# Paths
# -----------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
REPO_ROOT = SCRIPT_DIR.parent
ASSETS_DIR = SCRIPT_DIR / "assets"
OUTPUT_DOCX = SCRIPT_DIR / "CloudLens_Ansible_AWS_Customer_Runbook.docx"

# -----------------------------------------------------------------------------
# SVG -> PNG conversion (graceful fallback if cairosvg missing)
# -----------------------------------------------------------------------------
try:
    import cairosvg  # type: ignore

    HAS_CAIROSVG = True
except Exception:  # pragma: no cover
    HAS_CAIROSVG = False


def svg_to_png(svg_path: Path, png_path: Path, width: int = 1600) -> bool:
    if not HAS_CAIROSVG or not svg_path.exists():
        return False
    try:
        cairosvg.svg2png(url=str(svg_path), write_to=str(png_path), output_width=width)
        return True
    except Exception as exc:  # pragma: no cover
        print(f"[warn] cairosvg failed on {svg_path.name}: {exc}", file=sys.stderr)
        return False


# -----------------------------------------------------------------------------
# Low-level XML helpers
# -----------------------------------------------------------------------------
def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = "BFBFBF", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_MUTED


# Static, pre-evaluated Table of Contents. Kept in sync with build_* in main().
TOC_ENTRIES: list[tuple[int, str]] = [
    (1, "1. Executive Summary"),
    (1, "2. Solution Overview"),
    (1, "3. Choosing Your Deployment Path"),
    (1, "4. Prerequisites Checklist"),
    (1, "5. Deployment, Step by Step"),
    (2, "5.1  Tier 1: One-Click Launch Stack (CloudFormation)"),
    (2, "5.2  Tier 2: AWS CloudShell"),
    (2, "5.3  Tier 3: Docker (laptop or CI/CD)"),
    (1, "6. Supported EC2 Scenarios"),
    (1, "7. Verification Checklist"),
    (1, "8. Scaling Guide"),
    (1, "9. Troubleshooting Reference"),
    (1, "Appendix A: customer_input.yaml Schema"),
    (1, "Appendix B: Bulk Tag Script"),
    (1, "Appendix C: Quick Links"),
]


def _add_toc_entry(doc: Document, level: int, label: str) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    if level >= 2:
        pf.left_indent = Inches(0.35)
    tab_stops = pf.tab_stops
    try:
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    except Exception:
        tab_stops.add_tab_stop(Inches(6.3))
    label_run = p.add_run(label)
    label_run.font.name = "Calibri"
    label_run.font.size = Pt(12) if level == 1 else Pt(10.5)
    label_run.bold = (level == 1)
    label_run.font.color.rgb = AWS_BLUE if level == 1 else TEXT_DARK
    p.add_run("\t")


# -----------------------------------------------------------------------------
# Style configuration
# -----------------------------------------------------------------------------
def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_DARK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = AWS_NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = AWS_BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = KEYSIGHT_NAVY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True


# -----------------------------------------------------------------------------
# Reusable element builders
# -----------------------------------------------------------------------------
def add_paragraph(doc: Document, text: str, *, bold: bool = False,
                  italic: bool = False, color: RGBColor | None = None,
                  size: int = 11, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODE_BG)
    p_pr.append(shd)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = KEYSIGHT_NAVY


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _shade_cell(cell, CALLOUT_BG)
    _set_cell_borders(cell, color=ACCENT_HEX, size="6")
    label_p = cell.paragraphs[0]
    label_p.paragraph_format.space_after = Pt(2)
    run = label_p.add_run(label)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = AWS_NAVY
    body_p = cell.add_paragraph()
    run2 = body_p.add_run(body)
    run2.font.name = "Calibri"
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = TEXT_DARK
    doc.add_paragraph()


def add_checkbox_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.1)
        box = p.add_run("☐  ")
        box.font.name = "Segoe UI Symbol"
        box.font.size = Pt(13)
        box.font.color.rgb = AWS_BLUE
        text = p.add_run(item)
        text.font.name = "Calibri"
        text.font.size = Pt(11)
        text.font.color.rgb = TEXT_DARK


def add_styled_table(doc: Document, headers: list[str], rows: list[list[str]],
                     col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _shade_cell(cell, TABLE_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(header)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            cell = row_cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                _shade_cell(cell, TABLE_ALT_ROW_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_DARK
    if col_widths:
        for row in table.rows:
            for c_idx, width_in in enumerate(col_widths):
                row.cells[c_idx].width = Inches(width_in)
    doc.add_paragraph()


def add_image_or_placeholder(doc: Document, svg_name: str, caption: str,
                             width_inches: float = 6.5,
                             tmpdir: Path | None = None) -> None:
    svg_path = ASSETS_DIR / svg_name
    if HAS_CAIROSVG and tmpdir is not None and svg_path.exists():
        png_path = tmpdir / (svg_path.stem + ".png")
        if svg_to_png(svg_path, png_path, width=1600):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(png_path), width=Inches(width_inches))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap.add_run(caption)
            cap_run.italic = True
            cap_run.font.size = Pt(9.5)
            cap_run.font.color.rgb = TEXT_MUTED
            return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ Diagram: {svg_name} (see docs/assets/{svg_name}) ]")
    run.italic = True
    run.font.color.rgb = TEXT_MUTED
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9.5)
    cap_run.font.color.rgb = TEXT_MUTED


# -----------------------------------------------------------------------------
# Header / Footer
# -----------------------------------------------------------------------------
def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.different_first_page_header_footer = True

    header = section.header
    h_p = header.paragraphs[0]
    h_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_run = h_p.add_run("CloudLens Ansible for AWS: Customer Runbook")
    h_run.font.name = "Calibri"
    h_run.font.size = Pt(9)
    h_run.font.color.rgb = AWS_BLUE
    h_run.italic = True

    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = f_p.add_run("Keysight Technologies   |   ")
    left.font.name = "Calibri"
    left.font.size = Pt(9)
    left.font.color.rgb = TEXT_MUTED
    _add_page_number(f_p)


# -----------------------------------------------------------------------------
# Cover page
# -----------------------------------------------------------------------------
def build_cover_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_p.add_run("KEYSIGHT")
    logo_run.font.name = "Calibri"
    logo_run.font.size = Pt(20)
    logo_run.bold = True
    logo_run.font.color.rgb = KEYSIGHT_GOLD

    tagline_p = doc.add_paragraph()
    tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tagline_p.add_run("TECHNOLOGIES")
    tag_run.font.name = "Calibri"
    tag_run.font.size = Pt(10)
    tag_run.font.color.rgb = KEYSIGHT_NAVY
    tag_run.bold = True

    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CloudLens Ansible")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(36)
    title_run.bold = True
    title_run.font.color.rgb = AWS_NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("AWS Customer Runbook")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(26)
    sub_run.font.color.rgb = AWS_ORANGE

    doc.add_paragraph()
    doc.add_paragraph()

    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag_p.add_run(
        "Automated sensor deployment for Linux and Windows EC2 at scale"
    )
    tag_run.font.name = "Calibri"
    tag_run.font.size = Pt(14)
    tag_run.italic = True
    tag_run.font.color.rgb = TEXT_MUTED

    for _ in range(8):
        doc.add_paragraph()

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    for row in meta_table.rows:
        for cell in row.cells:
            cell.width = Inches(2.0)
    labels = [("Version", "v1.0.0"), ("Date", "June 2026")]
    for r_idx, (label, value) in enumerate(labels):
        l_cell = meta_table.rows[r_idx].cells[0]
        v_cell = meta_table.rows[r_idx].cells[1]
        _shade_cell(l_cell, TABLE_HEADER_BG)
        _shade_cell(v_cell, TABLE_ALT_ROW_BG)
        lp = l_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        lr.font.name = "Calibri"
        lr.font.size = Pt(11)
        vp = v_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run(value)
        vr.bold = True
        vr.font.color.rgb = KEYSIGHT_NAVY
        vr.font.name = "Calibri"
        vr.font.size = Pt(11)

    for _ in range(2):
        doc.add_paragraph()

    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = foot_p.add_run("Keysight Technologies   |   Network Visibility Solutions")
    foot_run.font.name = "Calibri"
    foot_run.font.size = Pt(10)
    foot_run.italic = True
    foot_run.font.color.rgb = TEXT_MUTED

    doc.add_page_break()


# -----------------------------------------------------------------------------
# Section builders
# -----------------------------------------------------------------------------
def build_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    for level, label in TOC_ENTRIES:
        _add_toc_entry(doc, level, label)
    doc.add_page_break()


def build_executive_summary(doc: Document) -> None:
    doc.add_heading("1. Executive Summary", level=1)
    add_paragraph(
        doc,
        "CloudLens Ansible for AWS is a fully automated kit that deploys "
        "Keysight CloudLens sensors to every tagged EC2 instance in an AWS "
        "account. It works on Linux and Windows, from a single instance to "
        "5,000+, without any manual per-host steps.",
    )
    add_paragraph(
        doc,
        "It is built for customer DevOps and network teams who already run "
        "Ansible and need network visibility instrumented quickly, and for "
        "Keysight Sales Engineers who need a repeatable, demoable proof-of-value "
        "asset. The Ansible layer connects over SSH, AWS Systems Manager (SSM), "
        "or WinRM, so it works even in accounts where one of those is locked "
        "down.",
    )

    doc.add_heading("Why it matters", level=2)
    for b in [
        "Eliminates the per-instance SSH / RDP grind. Tag an instance with "
        "cloudlens=yes and the kit handles the rest.",
        "Three entry points (CloudFormation Launch Stack, CloudShell curl, or "
        "Docker), so customers run it wherever they already work.",
        "Same playbook portable to Azure and GCP, so multi-cloud teams keep one "
        "workflow.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(b)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK

    doc.add_heading("Proof points", level=2)
    add_styled_table(
        doc,
        headers=["Scenario", "Result", "Time"],
        rows=[
            ["847-instance fleet (mixed Ubuntu/RHEL/Windows)", "100% sensor registration", "45 min"],
            ["Ubuntu 22.04 (SSH)", "Sensor running", "4 min"],
            ["Windows Server 2022 (SSM)", "Sensor running", "6 min"],
            ["vController registration", "All sensors Connected", "<1 min"],
        ],
        col_widths=[3.8, 2.0, 0.9],
    )
    add_paragraph(
        doc,
        "Numbers from production deployments and the deployment runbook. "
        "Reference account region: us-east-1.",
        italic=True, size=10, color=TEXT_MUTED,
    )
    doc.add_page_break()


def build_solution_overview(doc: Document, tmpdir: Path) -> None:
    doc.add_heading("2. Solution Overview", level=1)
    add_image_or_placeholder(
        doc, "architecture-diagram.svg",
        caption="Figure 1: End-to-end architecture (control point -> AWS inventory -> OS lanes -> vController)",
        width_inches=6.5, tmpdir=tmpdir,
    )
    add_paragraph(
        doc,
        "A single Ansible control point authenticates to AWS (SSO profile, "
        "access keys, or a CloudShell session), queries EC2 for every instance "
        "matching the customer's tag filter, and routes each host to the "
        "OS-specific playbook lane: Ubuntu, RHEL, or Windows.",
    )
    add_paragraph(
        doc,
        "Linux hosts run the CloudLens agent in a Docker or Podman container. "
        "Windows hosts run the CloudLens Windows sensor as a native service. "
        "Every sensor self-registers with CloudLens vController (formerly CLMS) "
        "on first start using the project key supplied in customer_input.yaml. "
        "For the packet path, KVO orchestrates AWS VPC Traffic Mirroring into "
        "collector SVMs and a vPB. No per-instance UI steps, no static inventory "
        "files to maintain.",
    )
    doc.add_page_break()


def build_choose_path(doc: Document, tmpdir: Path) -> None:
    doc.add_heading("3. Choosing Your Deployment Path", level=1)
    add_paragraph(
        doc,
        "All three paths run the same Ansible engine. Same playbooks, same "
        "automation. Pick the entry point that matches how your team works.",
    )
    add_image_or_placeholder(
        doc, "decision-tree.svg",
        caption="Figure 2: Decision tree (pick the entry point that matches your environment)",
        width_inches=6.0, tmpdir=tmpdir,
    )
    add_styled_table(
        doc,
        headers=["Tier", "Best For", "Tools Needed", "Effort"],
        rows=[
            ["Tier 1: Launch Stack",
             "Customers who live in the AWS Console and want zero local setup",
             "Web browser only",
             "Lowest"],
            ["Tier 2: CloudShell",
             "Customers already authenticated to AWS in the browser",
             "AWS CloudShell (pre-authenticated)",
             "Low"],
            ["Tier 3: Docker",
             "Repeatable runs from a laptop or CI/CD pipeline",
             "Docker, AWS credentials, customer_input.yaml",
             "Medium"],
        ],
        col_widths=[1.7, 2.5, 1.7, 0.9],
    )
    doc.add_page_break()


def build_prerequisites(doc: Document) -> None:
    doc.add_heading("4. Prerequisites Checklist (Printable)", level=1)
    add_paragraph(
        doc,
        "Confirm every item below before kicking off a deployment. Tick each "
        "box as you go. The kit will not magically fix a missing prerequisite.",
    )
    add_checkbox_list(doc, [
        "AWS account with rights to describe EC2 and (for Tier 1) create the "
        "CloudFormation stack.",
        "Subscribed to the three Keysight Marketplace AMIs (Vision One, CloudLens "
        "Manager, Virtual Packet Broker), one time per account.",
        "vController (CLMS) reachable from the target instance subnets on TCP/443.",
        "Project Key obtained from the vController UI (Settings > Projects > API Keys).",
        "Target instances tagged with cloudlens=yes, os=ubuntu|rhel|windows, env=prod|dev|qa.",
        "EC2 key pair (Linux SSH) or SSM Agent + IAM role (SSM mode) ready.",
    ])
    add_callout(
        doc,
        "Tip",
        "Tags are how the dynamic inventory discovers hosts. If an instance is "
        "not tagged, the kit cannot see it. That is the single most common "
        "support ticket. Use the bulk-tag script in Appendix B to tag an entire "
        "region in one shot.",
    )
    doc.add_page_break()


def build_deployment(doc: Document) -> None:
    doc.add_heading("5. Deployment, Step by Step", level=1)

    # ---- Tier 1 ----
    doc.add_heading("5.1  Tier 1: One-Click Launch Stack (CloudFormation)", level=2)
    add_paragraph(
        doc,
        "Best when you want zero local tools. The Launch Stack button opens the "
        "CloudFormation quick-create console with the stack template pre-loaded "
        "from this repository. It deploys vController, KVO (optional), and vPB "
        "(optional) from the Marketplace AMIs, then you run the sensor playbook.",
    )
    for step in [
        "Open the README or the public site and click the 'Launch Stack' button.",
        "Sign in to the AWS Console when prompted; the CloudFormation "
        "quick-create form opens with the template URL filled in.",
        "Fill the parameters: EC2 key pair, admin ingress CIDR, and which "
        "components to deploy (KVO and vPB default to yes).",
        "Acknowledge IAM capabilities and click Create Stack. About 5 minutes "
        "later the stack reaches CREATE_COMPLETE.",
        "Read the stack Outputs for the vController, KVO, and vPB URLs and IPs.",
        "Complete the vController EULA + first login, create a project, and copy "
        "the project key, then run the sensor playbook (Tier 2 or Tier 3).",
    ]:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
    add_callout(
        doc, "Expected outcome",
        "Stack reaches CREATE_COMPLETE in about 5 minutes. vController, KVO, and "
        "vPB come up on their own Elastic IPs. Instance types are fixed to the "
        "size each AMI is qualified on (vController and vPB t3.xlarge, KVO "
        "c5.2xlarge).",
    )

    # ---- Tier 2 ----
    doc.add_heading("5.2  Tier 2: AWS CloudShell", level=2)
    add_paragraph(
        doc,
        "Best when the customer is already logged into AWS in a browser tab. "
        "CloudShell is pre-authenticated, so there is no access key to manage "
        "and no local install.",
    )
    for step in [
        "Open AWS CloudShell from the console toolbar (region us-east-1).",
        "Run the quickstart bootstrap:",
        "Answer the wizard prompts: vController IP, project key, tag filter, "
        "connection mode (SSH / SSM / WinRM).",
        "Confirm sensors in vController > Sensors (about 8 min for a small fleet).",
    ]:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
    add_code_block(
        doc,
        "curl -sSL https://raw.githubusercontent.com/Keysight-Tech/"
        "cloudlens-ansible-aws/main/quickstart.sh | bash",
    )
    add_callout(
        doc, "Expected outcome",
        "Wizard discovers tagged instances, auto-tunes Ansible forks, deploys "
        "sensors, and prints a success summary. All state lives in your "
        "CloudShell home directory, with nothing installed locally.",
    )

    # ---- Tier 3 ----
    doc.add_heading("5.3  Tier 3: Docker (laptop or CI/CD)", level=2)
    add_paragraph(
        doc,
        "Best for repeatable runs from a developer laptop, a CI pipeline, or any "
        "container host. The image is pinned and hermetic, so you get the same "
        "result on macOS, Windows, Linux, GitHub Actions, GitLab CI, and Jenkins.",
    )
    for step in [
        "Copy customer_input.yaml.example to customer_input.yaml and fill in "
        "your vController IP, project key, region, and tag filter.",
        "Provide AWS credentials to the container (mount ~/.aws read-only, or "
        "pass AWS_ACCESS_KEY_ID / AWS_SECRET_ACCESS_KEY as env vars).",
        "For Linux SSH targets, mount your EC2 key pair path read-only.",
        "Run the container:",
    ]:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK
    add_code_block(
        doc,
        "docker run --rm -it \\\n"
        "  -v $(pwd)/customer_input.yaml:/work/customer_input.yaml \\\n"
        "  -v $HOME/.aws:/root/.aws:ro \\\n"
        "  -v $HOME/.ssh:/root/.ssh:ro \\\n"
        "  -e AWS_PROFILE -e AWS_DEFAULT_REGION=us-east-1 \\\n"
        "  ghcr.io/keysight-tech/cloudlens-ansible-aws:latest",
    )
    add_callout(
        doc, "Expected outcome",
        "Container exits 0 with a 'sensors deployed: N/N' summary. Pin the image "
        "tag in CI to get hermetic, reproducible runs.",
    )
    doc.add_page_break()


def build_scenarios(doc: Document, tmpdir: Path) -> None:
    doc.add_heading("6. Supported EC2 Scenarios", level=1)
    add_image_or_placeholder(
        doc, "scenario-matrix.svg",
        caption="Figure 3: EC2 compatibility matrix (OS x topology x connection method)",
        width_inches=6.5, tmpdir=tmpdir,
    )
    add_styled_table(
        doc,
        headers=["OS / Topology", "Public IP + SSH", "Private + Bastion",
                 "SSM (no inbound)", "CloudShell"],
        rows=[
            ["Ubuntu 20.04 / 22.04 / 24.04", "Supported", "Supported", "Supported", "Supported"],
            ["RHEL 7 / 8 / 9, Amazon Linux", "Supported", "Supported", "Supported", "Supported"],
            ["Rocky / AlmaLinux", "Supported", "Supported", "Supported", "Supported"],
            ["Windows Server 2019 / 2022", "Supported (WinRM)", "Supported (SSM)", "Supported", "Supported"],
        ],
        col_widths=[2.4, 1.3, 1.3, 1.3, 1.0],
    )
    doc.add_page_break()


def build_verification(doc: Document) -> None:
    doc.add_heading("7. Verification Checklist (Printable)", level=1)
    add_paragraph(
        doc,
        "After every deployment, walk this list. If any item fails, jump to "
        "section 9 (Troubleshooting Reference).",
    )
    add_checkbox_list(doc, [
        "All sensors visible in vController > Sensors page",
        "Filter by custom_tags (for example Customer=Acme) returns the expected instances",
        "KVO > Inventory > Devices shows vController and vPB as CONNECTED",
        "KVO > Cloud Fabric > Cloud Configs shows the AWS config COMMITTED",
        "Container logs show no errors (Linux): docker logs cloudlens-agent",
        "CloudLens service running (Windows): Get-Service CloudLensAgent",
    ])
    doc.add_page_break()


def build_scaling(doc: Document) -> None:
    doc.add_heading("8. Scaling Guide", level=1)
    add_styled_table(
        doc,
        headers=["Fleet Size", "Parallelism", "Sharded?", "Approx Time"],
        rows=[
            ["1 to 50", "50 concurrent", "No", "5 to 10 min"],
            ["51 to 200", "100 concurrent", "No", "10 to 20 min"],
            ["201 to 800", "200 concurrent", "No", "15 to 30 min"],
            ["801 to 2,000", "400 concurrent", "Yes (auto)", "30 to 60 min"],
            ["2,001 to 5,000", "800 concurrent", "Yes", "1 to 2 hr"],
            ["10,000+", "2,500+ concurrent", "Yes", "4+ hr"],
        ],
        col_widths=[1.4, 1.7, 1.4, 1.5],
    )
    add_paragraph(
        doc,
        "The kit auto-tunes the Ansible fork count based on the number of "
        "instances the dynamic inventory discovers. Above 2,000 instances it "
        "also auto-shards: the inventory is split into batches and the playbook "
        "runs them in controlled waves. This keeps memory and AWS API throttling "
        "under control while still finishing in a reasonable window for tens of "
        "thousands of hosts. See docs/SCALING.md in the repository for KVO "
        "infrastructure sizing and VPC Traffic Mirroring limits.",
    )
    doc.add_page_break()


def build_troubleshooting(doc: Document) -> None:
    doc.add_heading("9. Troubleshooting Reference", level=1)
    add_paragraph(
        doc,
        "Organised by symptom area. The full reference lives at "
        "docs/TROUBLESHOOTING.md and docs/OPERATIONS.md in the repository.",
    )

    doc.add_heading("Inventory", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely Cause", "Fix"],
        rows=[
            ["Inventory finds 0 instances",
             "Tags missing on instances",
             "aws ec2 create-tags --resources <id> --tags Key=cloudlens,Value=yes"],
            ["Subset of instances missing",
             "Tag value mismatch (case-sensitive)",
             "Use exact lowercase values: ubuntu, rhel, windows, prod"],
        ],
        col_widths=[2.2, 2.1, 2.7],
    )

    doc.add_heading("SSH (Linux)", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely Cause", "Fix"],
        rows=[
            ["Permission denied (publickey)",
             "Wrong key path or SSH user",
             "Set aws.ssh_key_path and the correct ssh_user_* in customer_input.yaml"],
            ["Connection timeout",
             "Security group blocks port 22 from control point",
             "Switch to bastion or SSM connection mode"],
        ],
        col_widths=[2.2, 2.1, 2.7],
    )

    doc.add_heading("SSM (Linux and Windows)", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely Cause", "Fix"],
        rows=[
            ["SSM shows 0 target instances",
             "Missing IAM role or SSM Agent stopped",
             "Attach AmazonSSMManagedInstanceCore; confirm aws ssm describe-instance-information"],
            ["session-manager-plugin not found",
             "AWS CLI SSM plugin not installed on control node",
             "Install session-manager-plugin, or switch linux_connection to ssh"],
        ],
        col_widths=[2.2, 2.1, 2.7],
    )

    doc.add_heading("Registration with vController", level=3)
    add_styled_table(
        doc,
        headers=["Symptom", "Likely Cause", "Fix"],
        rows=[
            ["Sensor not in vController UI",
             "Wrong project key",
             "Check Settings > Projects > API Keys; update customer_input.yaml"],
            ["Sensor offline after restart",
             "vController IP unreachable from instance subnet",
             "Open egress from the subnet to vController on TCP/443"],
        ],
        col_widths=[2.2, 2.1, 2.7],
    )
    doc.add_page_break()


def build_appendix_a(doc: Document) -> None:
    doc.add_heading("Appendix A: customer_input.yaml Schema", level=1)
    add_paragraph(
        doc,
        "Full annotated example. Copy customer_input.yaml.example from the "
        "repository to customer_input.yaml and fill in values. Never commit "
        "customer_input.yaml to git. Credentials belong in your AWS profile or "
        "env vars.",
    )
    add_code_block(doc, """# === AWS Environment ===
aws:
  profile: "default"
  regions:
    - us-east-1

  # Tag selector: instances are matched by these tags
  tag_filters:
    cloudlens: "yes"
    env:       "prod"

  windows_connection: "ssm"    # ssm | winrm
  linux_connection:   "ssh"    # ssh | ssm
  ssh_key_path:    "~/.aws/your-key.pem"
  ssh_user_ubuntu: "ubuntu"
  ssh_user_rhel:   "ec2-user"

# === CloudLens Configuration ===
cloudlens:
  manager_ip_or_fqdn: "vcontroller.customer.example.com"
  project_key:        "REPLACE_WITH_PROJECT_KEY"
  custom_tags:        "Env=AWS Region=us-east-1 Customer=Acme"
  registry_type:      "insecure"   # or "secure"
  linux_runtime:      "auto"       # auto | docker | podman

# Optional scoping: empty lists = all VPCs / subnets in region
vpc_ids:    []
subnet_ids: []

# === Deployment Behavior ===
deploy:
  forks: 0                # 0 = auto-tune
  timeout_seconds: 60
  reinstall_if_unhealthy: true
  shard_size: 500         # for >2000 instances
""")
    doc.add_page_break()


def build_appendix_b(doc: Document) -> None:
    doc.add_heading("Appendix B: Bulk Tag Script", level=1)
    add_paragraph(
        doc,
        "Apply the three required tags to instances in a region, filtered by "
        "platform. Run these from AWS CloudShell or any shell with the AWS CLI "
        "configured.",
    )

    doc.add_heading("Ubuntu / Linux instances", level=3)
    add_code_block(doc, """aws ec2 describe-instances --region us-east-1 \\
  --filters Name=instance-state-name,Values=running \\
            Name=platform-details,Values="Linux/UNIX" \\
  --query "Reservations[].Instances[].InstanceId" --output text \\
| xargs -n1 -I {} aws ec2 create-tags --resources {} \\
    --tags Key=cloudlens,Value=yes Key=os,Value=ubuntu Key=env,Value=prod""")

    doc.add_heading("Windows instances", level=3)
    add_code_block(doc, """aws ec2 describe-instances --region us-east-1 \\
  --filters Name=instance-state-name,Values=running \\
            Name=platform,Values=windows \\
  --query "Reservations[].Instances[].InstanceId" --output text \\
| xargs -n1 -I {} aws ec2 create-tags --resources {} \\
    --tags Key=cloudlens,Value=yes Key=os,Value=windows Key=env,Value=prod""")
    doc.add_page_break()


def build_appendix_c(doc: Document) -> None:
    doc.add_heading("Appendix C: Quick Links", level=1)
    add_styled_table(
        doc,
        headers=["Resource", "Location"],
        rows=[
            ["GitHub repository", "https://github.com/Keysight-Tech/cloudlens-ansible-aws"],
            ["Public site",
             "https://keysight-tech.github.io/cloudlens-ansible-aws/"],
            ["Deployment guide", "docs/DEPLOYMENT_GUIDE.md (in the repo)"],
            ["Operations guide", "docs/OPERATIONS.md (in the repo)"],
            ["Scaling guide", "docs/SCALING.md (in the repo)"],
            ["Troubleshooting guide", "docs/TROUBLESHOOTING.md (in the repo)"],
            ["Customer email templates", "docs/CUSTOMER_EMAIL.md (in the repo)"],
            ["GitHub Issues (bug reports / feature requests)",
             "https://github.com/Keysight-Tech/cloudlens-ansible-aws/issues"],
        ],
        col_widths=[2.6, 4.4],
    )
    add_paragraph(
        doc,
        "Support email subject (for the Keysight account team): "
        "'CloudLens Ansible AWS: <customer name>, <brief issue>'.",
        italic=True, size=10, color=TEXT_MUTED,
    )


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------
def main() -> int:
    if not HAS_CAIROSVG:
        print("[warn] cairosvg unavailable; diagrams will be rendered as captions.")
    print(f"[info] writing {OUTPUT_DOCX}")

    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)
    build_cover_page(doc)

    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        build_toc(doc)
        build_executive_summary(doc)
        build_solution_overview(doc, tmpdir)
        build_choose_path(doc, tmpdir)
        build_prerequisites(doc)
        build_deployment(doc)
        build_scenarios(doc, tmpdir)
        build_verification(doc)
        build_scaling(doc)
        build_troubleshooting(doc)
        build_appendix_a(doc)
        build_appendix_b(doc)
        build_appendix_c(doc)
        doc.save(OUTPUT_DOCX)

    size_kb = OUTPUT_DOCX.stat().st_size / 1024
    print(f"[ok]  wrote {OUTPUT_DOCX.name} ({size_kb:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
